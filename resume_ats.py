import streamlit as st
import fitz  # PyMuPDF
import docx2txt
import re
from docx import Document
from docx.shared import Pt
from io import BytesIO
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

st.set_page_config(page_title="ATS Resume Converter", layout="wide")

if "experiences" not in st.session_state:
    st.session_state.experiences = []
if "educations" not in st.session_state:
    st.session_state.educations = []
if "projects" not in st.session_state:
    st.session_state.projects = []
if "certifications" not in st.session_state:
    st.session_state.certifications = []
if "skills" not in st.session_state:
    st.session_state.skills = []

# Extract text from PDF

def extract_text_from_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

# Extract text from DOCX

def extract_text_from_docx(file):
    return docx2txt.process(file)

# Clean and normalize text

def clean_text(text):
    text = re.sub(r"\n{2,}", "\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()

# Parse fields from resume text

def parse_resume_sections(text):
    name = re.findall(r"(?m)^(?!.*@|\d)([A-Z][a-z]+\s[A-Z][a-z]+.*)", text)
    email = re.search(r"[\w\.-]+@[\w\.-]+", text)
    phone = re.search(r"\+?\d[\d\s\-]{8,}", text)
    educations = re.findall(r"(?i)([^\n]*University[^\n]*\n?.*?)\n(?=\S)", text)
    experiences = re.findall(r"(?i)([^\n]*Intern[^\n]*\n?.*?)\n(?=\S)", text)
    cert_keywords = ["coursera", "linkedin learning", "udemy", "certification", "certified"]
    certifications = [line.strip() for line in text.splitlines() if any(kw in line.lower() for kw in cert_keywords)]
    skills = [line.strip() for line in text.splitlines() if 'skills' in line.lower() or 'technologies' in line.lower()]

    lines = text.splitlines()
    projects = []
    current_block = []
    seen = set()

    def is_new_project_start(line):
        return any(keyword in line.lower() for keyword in ["connect", "discord", "gesture", "unet", "chatbot", "project", "description"]) and len(line.strip().split()) < 12

    for line in lines:
        if is_new_project_start(line) and current_block:
            block_text = "\n".join(current_block).strip()
            normalized = re.sub(r"\s+", " ", block_text.lower())[:100]
            if normalized not in seen:
                seen.add(normalized)
                title = current_block[0]
                description = "\n".join(current_block[1:])
                projects.append({"title": title.strip(), "description": description.strip()})
            current_block = []
        current_block.append(line)

    if current_block:
        block_text = "\n".join(current_block).strip()
        normalized = re.sub(r"\s+", " ", block_text.lower())[:100]
        if normalized not in seen:
            seen.add(normalized)
            title = current_block[0]
            description = "\n".join(current_block[1:])
            projects.append({"title": title.strip(), "description": description.strip()})

    return {
        "name": name[0] if name else "",
        "email": email.group() if email else "",
        "phone": phone.group() if phone else "",
        "educations": educations,
        "experiences": experiences,
        "projects": projects,
        "certifications": certifications,
        "skills": skills[:5]
    }

# ATS Score Function

def calculate_ats_score(resume_text, job_description):
    tfidf = TfidfVectorizer().fit_transform([resume_text, job_description])
    score = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0] * 100
    return round(score, 2)

# Generate final DOCX

def format_docx(sections):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for section, content in sections.items():
        if content:
            doc.add_heading(section, level=1)
            if section == "Projects":
                for proj in content:
                    doc.add_paragraph(proj["title"], style='List Bullet')
                    doc.add_paragraph(proj["description"])
            elif section in ["Certifications", "Skills"]:
                for line in content:
                    doc.add_paragraph(line.strip(), style='List Bullet')
            elif isinstance(content, list):
                for line in content:
                    doc.add_paragraph(line.strip())
            else:
                for line in content.strip().split("\n"):
                    doc.add_paragraph(line.strip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI
st.title("📄 ATS Resume Converter – Premium Features for Free")

uploaded_file = st.file_uploader("Upload Your Resume (PDF or DOCX)", type=["pdf", "docx"])
job_description = st.text_area("Paste Job Description Here (Optional for ATS Score)")

if uploaded_file:
    text = extract_text_from_pdf(uploaded_file) if uploaded_file.name.endswith(".pdf") else extract_text_from_docx(uploaded_file)
    text = clean_text(text)
    parsed = parse_resume_sections(text)

    ats_score = calculate_ats_score(text, job_description) if job_description else None

    st.subheader("🔍 Resume Summary")
    for key, val in parsed.items():
        if key == "projects":
            for proj in val:
                st.markdown(f"**{proj['title']}**\n{proj['description']}")
        elif isinstance(val, list):
            for item in val:
                st.markdown(f"- {item}")
        else:
            st.markdown(f"**{key.capitalize()}**: {val}")

    if ats_score:
        st.metric("🎯 ATS Match Score", f"{ats_score}%")

    st.download_button("📥 Download ATS Resume (DOCX)", data=format_docx(parsed), file_name="ATS_Resume.docx")
